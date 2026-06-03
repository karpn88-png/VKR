import os
import io
import json
import datetime
import re
import time
from io import BytesIO
from urllib.parse import quote
from zoneinfo import ZoneInfo

import requests
from fastapi import FastAPI, HTTPException, UploadFile, File, Form, Body
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse

from docx import Document as DocxDocument

from sqlalchemy import create_engine, Column, Integer, String, LargeBinary, DateTime, ForeignKey, Text, text as sql_text
from sqlalchemy.dialects.postgresql import JSONB
from sqlalchemy.orm import sessionmaker, declarative_base

from gigachat_client import CHECK_CRITERIA, generate_report, get_selected_criteria

# ============================================================
# DATABASE SETUP
# ============================================================

DATABASE_URL = os.getenv("DATABASE_URL")
AI_URL = os.getenv("AI_URL")
APP_TIMEZONE = ZoneInfo(os.getenv("APP_TIMEZONE", "Asia/Novosibirsk"))

engine = create_engine(DATABASE_URL)
SessionLocal = sessionmaker(bind=engine)
Base = declarative_base()


def attachment_content_disposition(filename: str) -> str:
    safe_name = (filename or "attachment").replace("/", "_").replace("\\", "_").replace('"', "_")
    ascii_name = re.sub(r"[^A-Za-z0-9._ -]", "_", safe_name).strip() or "attachment"
    encoded_name = quote(safe_name, safe="")
    return f'attachment; filename="{ascii_name}"; filename*=UTF-8\'\'{encoded_name}'


def as_utc_datetime(value: datetime.datetime | None) -> datetime.datetime | None:
    if value is None:
        return None
    if value.tzinfo is None:
        return value.replace(tzinfo=datetime.timezone.utc)
    return value.astimezone(datetime.timezone.utc)


def serialize_datetime(value: datetime.datetime | None) -> str | None:
    utc_value = as_utc_datetime(value)
    if utc_value is None:
        return None
    return utc_value.isoformat().replace("+00:00", "Z")


def format_local_datetime(value: datetime.datetime | None) -> str:
    utc_value = as_utc_datetime(value)
    if utc_value is None:
        return ""
    return utc_value.astimezone(APP_TIMEZONE).strftime("%d.%m.%Y %H:%M:%S")


def add_multiline_report(document: DocxDocument, text: str | None) -> None:
    lines = (text or "—").splitlines() or ["—"]
    for line in lines:
        document.add_paragraph(line.rstrip())


def format_numeric_vector(vector, values_per_line: int = 8) -> list[str]:
    if not isinstance(vector, list) or not vector:
        return ["Вектор не сохранён. Запустите проверку документа заново."]

    values = []
    for value in vector:
        try:
            values.append(f"{float(value):.6f}")
        except (TypeError, ValueError):
            values.append(str(value))

    return [
        ", ".join(values[index:index + values_per_line])
        for index in range(0, len(values), values_per_line)
    ]


class Document(Base):
    __tablename__ = "documents"

    id = Column(Integer, primary_key=True, index=True)
    filename = Column(String, nullable=False)
    content = Column(LargeBinary, nullable=False)
    uploaded_at = Column(DateTime, default=datetime.datetime.utcnow)


class AnalysisResult(Base):
    """
    Храним:
    - llm_report (текстовый отчёт)
    - analysis_json (НО: только 4 метрики, не весь JSON)
    """
    __tablename__ = "analysis_results"

    id = Column(Integer, primary_key=True, index=True)
    doc_id = Column(Integer, ForeignKey("documents.id", ondelete="CASCADE"), index=True, nullable=False)

    # Здесь сохраняем ТОЛЬКО:
    # total_words, unique_words, uniqueness, embedding_dim
    analysis_json = Column(JSONB, nullable=True)

    # Здесь сохраняем полный отчёт LLM
    llm_report = Column(Text, nullable=True)

    created_at = Column(DateTime, default=datetime.datetime.utcnow)


class WorkThreadState(Base):
    __tablename__ = "work_thread_state"

    student_id = Column(Integer, primary_key=True, index=True)
    status = Column(String, default="Не проверено", nullable=False)
    teacher_status = Column(String, default="Не проверено", nullable=False)
    norm_status = Column(String, default="Не проверено", nullable=False)
    teacher_preliminary_grade = Column(String, default="", nullable=False)
    teacher_predefense_grade = Column(String, default="", nullable=False)
    teacher_final_grade = Column(String, default="", nullable=False)
    norm_preliminary_grade = Column(String, default="", nullable=False)
    norm_predefense_grade = Column(String, default="", nullable=False)
    norm_final_grade = Column(String, default="", nullable=False)
    submitted_at = Column(DateTime, nullable=True)
    checked_at = Column(DateTime, nullable=True)
    updated_at = Column(DateTime, default=datetime.datetime.utcnow, onupdate=datetime.datetime.utcnow)


class WorkMessage(Base):
    __tablename__ = "work_messages"

    id = Column(Integer, primary_key=True, index=True)
    student_id = Column(Integer, index=True, nullable=False)
    sender_role = Column(String, nullable=False)
    sender_name = Column(String, nullable=False)
    recipient_name = Column(String, nullable=True)
    text = Column(Text, nullable=True)
    message_type = Column(String, default="message", nullable=False)
    file_name = Column(String, nullable=True)
    file_content = Column(LargeBinary, nullable=True)
    file_content_type = Column(String, nullable=True)
    created_at = Column(DateTime, default=datetime.datetime.utcnow)


Base.metadata.create_all(bind=engine)


def ensure_work_thread_state_columns() -> None:
    """Add review columns for deployments created before separated statuses."""
    statements = [
        "ALTER TABLE work_thread_state ADD COLUMN IF NOT EXISTS teacher_status VARCHAR NOT NULL DEFAULT 'Не проверено'",
        "ALTER TABLE work_thread_state ADD COLUMN IF NOT EXISTS norm_status VARCHAR NOT NULL DEFAULT 'Не проверено'",
        "ALTER TABLE work_thread_state ADD COLUMN IF NOT EXISTS teacher_preliminary_grade VARCHAR NOT NULL DEFAULT ''",
        "ALTER TABLE work_thread_state ADD COLUMN IF NOT EXISTS teacher_predefense_grade VARCHAR NOT NULL DEFAULT ''",
        "ALTER TABLE work_thread_state ADD COLUMN IF NOT EXISTS teacher_final_grade VARCHAR NOT NULL DEFAULT ''",
        "ALTER TABLE work_thread_state ADD COLUMN IF NOT EXISTS norm_preliminary_grade VARCHAR NOT NULL DEFAULT ''",
        "ALTER TABLE work_thread_state ADD COLUMN IF NOT EXISTS norm_predefense_grade VARCHAR NOT NULL DEFAULT ''",
        "ALTER TABLE work_thread_state ADD COLUMN IF NOT EXISTS norm_final_grade VARCHAR NOT NULL DEFAULT ''",
        "UPDATE work_thread_state SET teacher_status = status WHERE teacher_status = 'Не проверено' AND status <> 'Не проверено'",
        "UPDATE work_thread_state SET norm_status = status WHERE norm_status = 'Не проверено' AND status <> 'Не проверено'",
    ]

    with engine.begin() as connection:
        for statement in statements:
            connection.execute(sql_text(statement))


ensure_work_thread_state_columns()

STUDENT_PROFILE = {
    "id": 1,
    "fio": "Иванов Иван Иванович",
    "short_name": "Иванов И. И.",
    "group": "АТ-23",
    "topic": "Разработка информационной системы",
    "teacher": "Бакаев Максим Александрович",
    "teacher_short": "Бакаев М. А.",
    "default_status": "Не проверено",
    "grade": "",
    "teacherGrade": "",
}

STUDENT_PROFILES = [
    STUDENT_PROFILE,
    {
        "id": 2,
        "fio": "Сидорова Анна Петровна",
        "short_name": "Сидорова А. П.",
        "group": "АТ-24",
        "topic": "Разработка мобильного приложения",
        "teacher": "Бакаев Максим Александрович",
        "teacher_short": "Бакаев М. А.",
        "default_status": "На проверке",
        "grade": "",
        "teacherGrade": "",
        "demo": True,
    },
    {
        "id": 3,
        "fio": "Петров Алексей Дмитриевич",
        "short_name": "Петров А. Д.",
        "group": "АО-22",
        "topic": "Разработка базы данных",
        "teacher": "Бакаев Максим Александрович",
        "teacher_short": "Бакаев М. А.",
        "default_status": "Требуется доработка",
        "grade": "",
        "teacherGrade": "",
        "demo": True,
    },
    {
        "id": 4,
        "fio": "Карпенко Никита Денисович",
        "short_name": "Карпенко Н. Д.",
        "group": "АТ-23",
        "topic": "Разработка информационной системы на основе нейросетевой модели для подготовки ВКР",
        "teacher": "Тетерин Максим Михайлович",
        "teacher_short": "Тетерин М. М.",
        "default_status": "Проверено",
        "default_teacher_status": "Проверено",
        "default_norm_status": "Требуется доработка",
        "grade": "",
        "teacherGrade": "",
        "demo": True,
    },
    {
        "id": 5,
        "fio": "Филатова Виктория Сергеевна",
        "short_name": "Филатова В. С.",
        "group": "АТ-23",
        "topic": "Разработка информационной системы для взаимодействия студентов и преподавателей при работе с ВКР",
        "teacher": "Тетерин Максим Михайлович",
        "teacher_short": "Тетерин М. М.",
        "default_status": "Проверено",
        "default_teacher_status": "Проверено",
        "default_norm_status": "Требуется доработка",
        "grade": "",
        "teacherGrade": "",
        "demo": True,
    },
]

TEACHER_PROFILE = {
    "short_name": "Тетерин М. М.",
    "full_name": "Тетерин Максим Михайлович",
}

NORMCONTROL_PROFILE = {
    "short_name": "Герасимов А. К.",
    "full_name": "Герасимов Антон Константинович",
}


def normalize_work_role(role: str | None) -> str:
    value = (role or "").strip().lower()
    if value in {"teacher", "supervisor", "преподаватель", "руководитель"}:
        return "teacher"
    if value in {"normcontrol", "norm", "reviewer", "нормоконтроль"}:
        return "normcontrol"
    if value == "student":
        return "student"
    return value


def recipient_profile_for_role(role: str | None) -> dict:
    normalized = normalize_work_role(role)
    if normalized == "normcontrol":
        return NORMCONTROL_PROFILE
    return TEACHER_PROFILE


def status_actor_name(role: str | None) -> str:
    normalized = normalize_work_role(role)
    if normalized == "normcontrol":
        return NORMCONTROL_PROFILE["short_name"]
    if normalized == "teacher":
        return TEACHER_PROFILE["short_name"]
    return STUDENT_PROFILE["short_name"]


def profile_teacher_short(profile: dict) -> str:
    return profile.get("teacher_short") or TEACHER_PROFILE["short_name"]


def profile_belongs_to_current_teacher(profile: dict) -> bool:
    return (
        profile_teacher_short(profile) == TEACHER_PROFILE["short_name"]
        or profile.get("teacher") == TEACHER_PROFILE["full_name"]
    )


def message_belongs_to_role(message: WorkMessage, role: str | None) -> bool:
    normalized = normalize_work_role(role)
    if normalized not in {"teacher", "normcontrol"}:
        return True

    norm_names = {
        NORMCONTROL_PROFILE["short_name"],
        NORMCONTROL_PROFILE["full_name"],
    }
    recipient_name = message.recipient_name or ""

    if normalized == "normcontrol":
        return message.sender_role == "normcontrol" or (
            message.sender_role == "student" and recipient_name in norm_names
        )

    return message.sender_role == "teacher" or (
        message.sender_role == "student" and recipient_name not in norm_names
    )


def get_student_profile(student_id: int) -> dict:
    return next(
        (profile for profile in STUDENT_PROFILES if profile["id"] == student_id),
        STUDENT_PROFILE,
    )


WORK_STATUSES = {
    "Не проверено",
    "На проверке",
    "Требуется доработка",
    "Проверено",
}

GRADE_KEY_ALIASES = {
    "preliminary": "preliminary",
    "preliminaryGrade": "preliminary",
    "predefense": "predefense",
    "predefenseGrade": "predefense",
    "defense": "predefense",
    "final": "final",
    "finalGrade": "final",
}

GRADE_COLUMNS = {
    "teacher": {
        "preliminary": "teacher_preliminary_grade",
        "predefense": "teacher_predefense_grade",
        "final": "teacher_final_grade",
    },
    "normcontrol": {
        "preliminary": "norm_preliminary_grade",
        "predefense": "norm_predefense_grade",
        "final": "norm_final_grade",
    },
}


def status_for_role(state: WorkThreadState, role: str | None) -> str:
    normalized = normalize_work_role(role)
    if normalized == "normcontrol":
        return state.norm_status
    return state.teacher_status


def set_status_for_role(state: WorkThreadState, role: str | None, status: str) -> None:
    normalized = normalize_work_role(role)
    if normalized == "normcontrol":
        state.norm_status = status
        return

    state.teacher_status = status
    state.status = status


def grades_for_role(state: WorkThreadState, role: str) -> dict:
    columns = GRADE_COLUMNS[role]
    return {
        "preliminary": getattr(state, columns["preliminary"]) or "",
        "defense": getattr(state, columns["predefense"]) or "",
        "final": getattr(state, columns["final"]) or "",
    }


def apply_grades_for_role(state: WorkThreadState, role: str, grades: dict) -> None:
    columns = GRADE_COLUMNS[role]
    for key, value in grades.items():
        normalized_key = GRADE_KEY_ALIASES.get(str(key))
        if not normalized_key:
            continue

        setattr(
            state,
            columns[normalized_key],
            "" if value is None else str(value).strip(),
        )


def teacher_flat_grades(state: WorkThreadState) -> dict:
    teacher_grades = grades_for_role(state, "teacher")
    return {
        "preliminaryGrade": teacher_grades["preliminary"],
        "predefenseGrade": teacher_grades["defense"],
        "finalGrade": teacher_grades["final"],
        "grade": teacher_grades["final"],
        "teacherGrade": teacher_grades["final"],
    }


def extract_rtf_text(content: bytes) -> str:
    """Best-effort RTF to plain text extraction without external binaries."""
    raw = content.decode("utf-8", errors="ignore")
    if not raw.strip():
        raw = content.decode("cp1251", errors="ignore")

    def hex_to_char(match):
        return bytes.fromhex(match.group(1)).decode("cp1251", errors="ignore")

    def unicode_to_char(match):
        code = int(match.group(1))
        if code < 0:
            code += 65536
        return chr(code)

    text = re.sub(r"\\'([0-9a-fA-F]{2})", hex_to_char, raw)
    text = re.sub(r"\\u(-?\d+).", unicode_to_char, text)
    text = re.sub(r"\\par[d]?", "\n", text)
    text = re.sub(r"\\[a-zA-Z]+-?\d* ?", " ", text)
    text = re.sub(r"\\[^a-zA-Z0-9]", " ", text)
    text = text.replace("{", " ").replace("}", " ")
    return re.sub(r"\s+", " ", text).strip()


def extract_legacy_doc_text(content: bytes) -> str:
    """Best-effort extraction for legacy .doc when no office converter exists."""
    decoded = content.decode("cp1251", errors="ignore")
    decoded = re.sub(r"[^\w\s.,:;!?()\-—«»\"'№%/\\]+", " ", decoded, flags=re.UNICODE)
    text = re.sub(r"\s+", " ", decoded).strip()

    if len(text) < 100:
        raise HTTPException(
            status_code=400,
            detail="Не удалось извлечь текст из .doc. Сохраните работу как .docx или .pdf и повторите проверку.",
        )

    return text


def run_ai_analysis(text: str) -> dict:
    if not AI_URL:
        raise HTTPException(status_code=503, detail="AI_URL is not configured")

    attempts = int(os.getenv("AI_REQUEST_ATTEMPTS", "6"))
    last_error = None

    for attempt in range(1, attempts + 1):
        try:
            response = requests.post(f"{AI_URL}/analyze", json={"text": text}, timeout=(5, 120))
            response.raise_for_status()
            return response.json()
        except requests.RequestException as e:
            last_error = e
            if attempt < attempts:
                time.sleep(min(2 * attempt, 10))

    raise HTTPException(
        status_code=503,
        detail=f"AI module unavailable after {attempts} attempts: {last_error}",
    )


def get_or_create_work_state(
    db,
    student_id: int,
    default_status: str = "Не проверено",
    default_teacher_status: str | None = None,
    default_norm_status: str | None = None,
) -> WorkThreadState:
    state = db.query(WorkThreadState).filter(WorkThreadState.student_id == student_id).first()
    if state:
        return state

    teacher_status = default_teacher_status or default_status
    norm_status = default_norm_status or default_status
    state = WorkThreadState(
        student_id=student_id,
        status=teacher_status,
        teacher_status=teacher_status,
        norm_status=norm_status,
    )
    db.add(state)
    db.commit()
    db.refresh(state)
    return state


def serialize_work_message(message: WorkMessage) -> dict:
    return {
        "id": message.id,
        "student_id": message.student_id,
        "sender_role": message.sender_role,
        "sender_name": message.sender_name,
        "recipient_name": message.recipient_name,
        "text": message.text,
        "message_type": message.message_type,
        "file_name": message.file_name,
        "has_file": bool(message.file_content),
        "file_content_type": message.file_content_type,
        "download_url": f"/work_thread/{message.student_id}/attachments/{message.id}" if message.file_content else None,
        "created_at": serialize_datetime(message.created_at),
    }


def serialize_work_thread(db, student_id: int, recipient_role: str | None = None) -> dict:
    student_profile = get_student_profile(student_id)
    state = get_or_create_work_state(
        db,
        student_id,
        student_profile.get("default_status", "Не проверено"),
        student_profile.get("default_teacher_status"),
        student_profile.get("default_norm_status"),
    )
    all_messages = (
        db.query(WorkMessage)
        .filter(WorkMessage.student_id == student_id)
        .order_by(WorkMessage.created_at.asc(), WorkMessage.id.asc())
        .all()
    )
    messages = [
        message
        for message in all_messages
        if message_belongs_to_role(message, recipient_role)
    ]

    return {
        "student": student_profile,
        "teacher": TEACHER_PROFILE,
        "normcontrol": NORMCONTROL_PROFILE,
        "status": status_for_role(state, recipient_role),
        "teacherStatus": state.teacher_status,
        "normStatus": state.norm_status,
        "teacherGrades": grades_for_role(state, "teacher"),
        "normGrades": grades_for_role(state, "normcontrol"),
        **teacher_flat_grades(state),
        "submitted_at": serialize_datetime(state.submitted_at),
        "checked_at": serialize_datetime(state.checked_at),
        "messages": [serialize_work_message(message) for message in messages],
    }

# ============================================================
# FASTAPI APP
# ============================================================

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# ============================================================
# ROUTES
# ============================================================


@app.get("/teacher_students")
def list_teacher_students(recipient_role: str = "teacher"):
    recipient_role = normalize_work_role(recipient_role) or "teacher"
    db = SessionLocal()
    try:
        students = []
        profiles = STUDENT_PROFILES
        if recipient_role == "teacher":
            profiles = [
                profile
                for profile in STUDENT_PROFILES
                if profile_belongs_to_current_teacher(profile)
            ]

        for profile in profiles:
            state = get_or_create_work_state(
                db,
                profile["id"],
                profile.get("default_status", "Не проверено"),
                profile.get("default_teacher_status"),
                profile.get("default_norm_status"),
            )
            students.append({
                "id": profile["id"],
                "fio": profile["fio"],
                "group": profile["group"],
                "topic": profile["topic"],
                "teacher": profile_teacher_short(profile),
                "status": status_for_role(state, recipient_role),
                "teacherStatus": state.teacher_status,
                "normStatus": state.norm_status,
                "teacherGrades": grades_for_role(state, "teacher"),
                "normGrades": grades_for_role(state, "normcontrol"),
                **teacher_flat_grades(state),
                "demo": bool(profile.get("demo")),
            })
        return students
    finally:
        db.close()


@app.get("/work_thread/{student_id}")
def get_work_thread(student_id: int, recipient_role: str | None = None):
    db = SessionLocal()
    try:
        return serialize_work_thread(db, student_id, recipient_role)
    finally:
        db.close()


@app.post("/work_thread/{student_id}/messages")
async def create_work_message(
    student_id: int,
    sender_role: str = Form(...),
    sender_name: str = Form(...),
    recipient_name: str = Form(default=""),
    recipient_role: str = Form(default=""),
    text: str = Form(default=""),
    message_type: str = Form(default="message"),
    file: UploadFile | None = File(default=None),
):
    sender_role = normalize_work_role(sender_role)
    recipient_role = normalize_work_role(recipient_role)
    if sender_role not in {"student", "teacher", "normcontrol"}:
        raise HTTPException(status_code=400, detail="sender_role must be student, teacher or normcontrol")

    text = text.strip()
    file_content = None
    file_name = None
    file_content_type = None

    if file and file.filename:
        file_content = await file.read()
        file_name = file.filename
        file_content_type = file.content_type

    if not text and not file_content:
        raise HTTPException(status_code=400, detail="Message text or file is required")

    db = SessionLocal()
    try:
        student_profile = get_student_profile(student_id)
        state = get_or_create_work_state(
            db,
            student_id,
            student_profile.get("default_status", "Не проверено"),
            student_profile.get("default_teacher_status"),
            student_profile.get("default_norm_status"),
        )
        if sender_role == "student" and file_content and recipient_role in {"teacher", "normcontrol"}:
            set_status_for_role(state, recipient_role, "На проверке")
            state.submitted_at = datetime.datetime.utcnow()
            state.checked_at = None

        message = WorkMessage(
            student_id=student_id,
            sender_role=sender_role,
            sender_name=sender_name.strip(),
            recipient_name=recipient_name.strip() or None,
            text=text or None,
            message_type=message_type.strip() or "message",
            file_name=file_name,
            file_content=file_content,
            file_content_type=file_content_type,
        )
        db.add(message)
        db.commit()
        db.refresh(message)
        return serialize_work_message(message)
    finally:
        db.close()


@app.get("/work_thread/{student_id}/attachments/{message_id}")
def download_work_attachment(student_id: int, message_id: int):
    db = SessionLocal()
    message = (
        db.query(WorkMessage)
        .filter(WorkMessage.student_id == student_id, WorkMessage.id == message_id)
        .first()
    )
    db.close()

    if not message or not message.file_content:
        raise HTTPException(status_code=404, detail="Attachment not found")

    safe_name = (message.file_name or f"attachment_{message_id}").replace("/", "_").replace("\\", "_")
    return StreamingResponse(
        BytesIO(message.file_content),
        media_type=message.file_content_type or "application/octet-stream",
        headers={"Content-Disposition": attachment_content_disposition(safe_name)},
    )


@app.post("/work_thread/{student_id}/submit")
async def submit_work(
    student_id: int,
    sender_name: str = Form(default=STUDENT_PROFILE["short_name"]),
    recipient_role: str = Form(default="teacher"),
    text: str = Form(default="Работа отправлена на проверку."),
    file: UploadFile | None = File(default=None),
):
    recipient = recipient_profile_for_role(recipient_role)
    recipient_role = normalize_work_role(recipient_role)

    db = SessionLocal()
    try:
        student_profile = get_student_profile(student_id)
        state = get_or_create_work_state(
            db,
            student_id,
            student_profile.get("default_status", "Не проверено"),
            student_profile.get("default_teacher_status"),
            student_profile.get("default_norm_status"),
        )
        set_status_for_role(state, recipient_role, "На проверке")
        state.submitted_at = datetime.datetime.utcnow()
        state.checked_at = None

        file_content = None
        file_name = None
        file_content_type = None
        if file and file.filename:
            file_content = await file.read()
            file_name = file.filename
            file_content_type = file.content_type

        message = WorkMessage(
            student_id=student_id,
            sender_role="student",
            sender_name=sender_name.strip() or STUDENT_PROFILE["short_name"],
            recipient_name=recipient["full_name"],
            text=text.strip() or "Работа отправлена на проверку.",
            message_type="submission",
            file_name=file_name,
            file_content=file_content,
            file_content_type=file_content_type,
        )
        db.add(message)
        db.commit()
        return serialize_work_thread(db, student_id, recipient_role)
    finally:
        db.close()


@app.post("/work_thread/{student_id}/mark_checked")
def mark_work_checked(student_id: int, checker_role: str = Form(default="teacher")):
    checker_role = normalize_work_role(checker_role)
    checker = recipient_profile_for_role(checker_role)
    student_profile = get_student_profile(student_id)

    db = SessionLocal()
    try:
        state = get_or_create_work_state(
            db,
            student_id,
            student_profile.get("default_status", "Не проверено"),
            student_profile.get("default_teacher_status"),
            student_profile.get("default_norm_status"),
        )
        set_status_for_role(state, checker_role, "Проверено")
        state.checked_at = datetime.datetime.utcnow()

        message = WorkMessage(
            student_id=student_id,
            sender_role=checker_role,
            sender_name=checker["short_name"],
            recipient_name=student_profile["fio"],
            text='Работа отмечена как "Проверено".',
            message_type="status",
        )
        db.add(message)
        db.commit()
        return serialize_work_thread(db, student_id, checker_role)
    finally:
        db.close()


@app.post("/work_thread/{student_id}/status")
def update_work_status(
    student_id: int,
    status: str = Form(...),
    actor_role: str = Form(default="teacher"),
):
    status = status.strip()
    if status not in WORK_STATUSES:
        raise HTTPException(status_code=400, detail=f"Unsupported work status: {status}")

    actor_role = normalize_work_role(actor_role)
    if actor_role not in {"teacher", "normcontrol"}:
        raise HTTPException(status_code=403, detail="Only teacher or normcontrol can change work status")

    student_profile = get_student_profile(student_id)
    db = SessionLocal()
    try:
        state = get_or_create_work_state(
            db,
            student_id,
            student_profile.get("default_status", "Не проверено"),
            student_profile.get("default_teacher_status"),
            student_profile.get("default_norm_status"),
        )
        set_status_for_role(state, actor_role, status)
        state.updated_at = datetime.datetime.utcnow()

        if status == "На проверке":
            state.submitted_at = state.submitted_at or datetime.datetime.utcnow()
            state.checked_at = None
        elif status == "Проверено":
            state.checked_at = datetime.datetime.utcnow()
        else:
            state.checked_at = None

        message = WorkMessage(
            student_id=student_id,
            sender_role=actor_role,
            sender_name=status_actor_name(actor_role),
            recipient_name=student_profile["fio"],
            text=f'Статус работы изменён на "{status}".',
            message_type="status",
        )
        db.add(message)
        db.commit()
        return serialize_work_thread(db, student_id, actor_role)
    finally:
        db.close()


@app.post("/work_thread/{student_id}/grades")
def update_work_grades(
    student_id: int,
    payload: dict | None = Body(default=None),
):
    payload = payload or {}
    actor_role = normalize_work_role(payload.get("actor_role") or "teacher")
    target_role = normalize_work_role(payload.get("target_role") or actor_role)

    if actor_role not in {"teacher", "normcontrol"}:
        raise HTTPException(status_code=403, detail="Only teacher or normcontrol can change grades")

    if target_role not in {"teacher", "normcontrol"}:
        raise HTTPException(status_code=400, detail="target_role must be teacher or normcontrol")

    if actor_role != target_role:
        raise HTTPException(status_code=403, detail="Users can change only their own role grades")

    grades = payload.get("grades")
    if not isinstance(grades, dict):
        raise HTTPException(status_code=400, detail="grades must be an object")

    db = SessionLocal()
    try:
        student_profile = get_student_profile(student_id)
        state = get_or_create_work_state(
            db,
            student_id,
            student_profile.get("default_status", "Не проверено"),
            student_profile.get("default_teacher_status"),
            student_profile.get("default_norm_status"),
        )
        apply_grades_for_role(state, target_role, grades)
        state.updated_at = datetime.datetime.utcnow()
        db.commit()
        return serialize_work_thread(db, student_id, actor_role)
    finally:
        db.close()


@app.post("/upload")
async def upload_file(file: UploadFile = File(...)):
    """Загрузка файла в базу данных"""
    data = await file.read()

    db = SessionLocal()
    doc = Document(filename=file.filename, content=data)
    db.add(doc)
    db.commit()
    db.refresh(doc)
    db.close()

    return {"status": "ok", "id": doc.id, "filename": file.filename}


@app.get("/documents")
def list_documents():
    """Вернуть список всех документов"""
    db = SessionLocal()
    docs = db.query(Document).order_by(Document.uploaded_at.desc()).all()
    db.close()

    return [
        {
            "id": d.id,
            "filename": d.filename,
            "uploaded_at": serialize_datetime(d.uploaded_at),
            "size": len(d.content),
        }
        for d in docs
    ]


@app.get("/document/{doc_id}")
def get_document(doc_id: int):
    """Получить информацию о документе"""
    db = SessionLocal()
    doc = db.query(Document).filter(Document.id == doc_id).first()
    db.close()

    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    return {
        "id": doc.id,
        "filename": doc.filename,
        "uploaded_at": serialize_datetime(doc.uploaded_at),
        "size": len(doc.content),
    }


@app.get("/analysis_criteria")
def analysis_criteria():
    """Вернуть пункты промпта, доступные для выбора перед LLM-проверкой."""
    return [
        {
            "id": item["id"],
            "title": item["title"],
            "instruction": item["instruction"],
        }
        for item in CHECK_CRITERIA
    ]


@app.post("/analyze_document/{doc_id}")
def analyze_document(doc_id: int, payload: dict | None = Body(default=None)):
    """Извлечь текст из файла, отправить в AI-модуль, сгенерировать LLM-отчёт и сохранить его в БД"""
    selected_checks = []
    if isinstance(payload, dict) and isinstance(payload.get("selected_checks"), list):
        selected_checks = [
            str(item)
            for item in payload.get("selected_checks", [])
            if isinstance(item, str)
        ]
    selected_criteria = get_selected_criteria(selected_checks)

    db = SessionLocal()
    doc = db.query(Document).filter(Document.id == doc_id).first()
    db.close()

    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    filename = doc.filename.lower()

    # ============================================================
    # ИЗВЛЕЧЕНИЕ ТЕКСТА
    # ============================================================

    if filename.endswith(".txt"):
        text = doc.content.decode("utf-8", errors="ignore")

    elif filename.endswith(".docx"):
        f = DocxDocument(io.BytesIO(doc.content))
        text = "\n".join(p.text for p in f.paragraphs)

    elif filename.endswith(".pdf"):
        import PyPDF2
        pdf = PyPDF2.PdfReader(io.BytesIO(doc.content))
        text = "\n".join(page.extract_text() or "" for page in pdf.pages)

    elif filename.endswith(".rtf"):
        text = extract_rtf_text(doc.content)

    elif filename.endswith(".doc"):
        text = extract_legacy_doc_text(doc.content)

    else:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file format: {filename}. Use .docx, .doc, .pdf, .rtf or .txt.",
        )

    if not text.strip():
        raise HTTPException(status_code=400, detail="Не удалось извлечь текст из файла.")

    # ============================================================
    # ОТПРАВКА В AI-МОДУЛЬ
    # ============================================================

    analysis = run_ai_analysis(text)

    embedding_vector = analysis.get("embedding") or []

    # Сохраняем метрики и числовой вектор текста для Word-отчета.
    brief = {
        "total_words": analysis.get("total_words"),
        "unique_words": analysis.get("unique_words"),
        "uniqueness": analysis.get("uniqueness"),
        "embedding_dim": analysis.get("embedding_dim"),
        "embedding_source": analysis.get("embedding_source"),
        "embedding_model": analysis.get("embedding_model"),
        "embedding_error": analysis.get("embedding_error"),
        "embedding_vector": embedding_vector,
        "selected_checks": [item["id"] for item in selected_criteria],
        "selected_check_titles": [item["title"] for item in selected_criteria],
    }
    llm_signals = {
        "total_words": brief["total_words"],
        "unique_words": brief["unique_words"],
        "uniqueness": brief["uniqueness"],
        "embedding_dim": brief["embedding_dim"],
        "embedding_source": brief["embedding_source"],
        "embedding_model": brief["embedding_model"],
        "selected_checks": brief["selected_check_titles"],
    }

    # ============================================================
    # GigaChat — отдельно, чтобы не ломать основной анализ
    # ============================================================

    try:
        llm_report = generate_report(
            text,
            local_signals={"local_analysis_summary": llm_signals},
            selected_checks=brief["selected_checks"],
        )
    except Exception as e:
        llm_report = f"[GigaChat error] {e}"

    # ============================================================
    # СОХРАНЯЕМ В БД
    # ============================================================

    db = SessionLocal()
    db.add(AnalysisResult(
        doc_id=doc_id,
        analysis_json=brief,
        llm_report=llm_report
    ))
    db.commit()
    db.close()

    return {
        "filename": filename,
        "analysis": llm_signals,
        "selected_checks": brief["selected_checks"],
        "llm_report": llm_report
    }


@app.get("/report_word/{doc_id}")
def report_word(doc_id: int):
    """Скачать последний LLM-отчёт в Word + 4 метрики"""
    db = SessionLocal()

    doc = db.query(Document).filter(Document.id == doc_id).first()
    if not doc:
        db.close()
        raise HTTPException(status_code=404, detail="Document not found")

    result = (
        db.query(AnalysisResult)
        .filter(AnalysisResult.doc_id == doc_id)
        .order_by(AnalysisResult.created_at.desc())
        .first()
    )
    db.close()

    if not result:
        raise HTTPException(status_code=404, detail="No analysis found for this document yet. Run analyze first.")

    d = DocxDocument()
    d.add_heading("Отчёт анализа ВКР", level=1)
    d.add_paragraph(f"Файл: {doc.filename}")
    d.add_paragraph(f"Дата: {format_local_datetime(result.created_at)}")

    a = result.analysis_json or {}

    d.add_heading("Выбранные пункты LLM-проверки", level=2)
    selected_titles = a.get("selected_check_titles") or []
    if selected_titles:
        for title in selected_titles:
            d.add_paragraph(str(title), style="List Bullet")
    else:
        d.add_paragraph("Использован полный набор критериев проверки.")

    d.add_heading("Метрики текста", level=2)
    d.add_paragraph(f"total_words: {a.get('total_words')}")
    d.add_paragraph(f"unique_words: {a.get('unique_words')}")
    d.add_paragraph(f"uniqueness: {a.get('uniqueness')}")
    d.add_paragraph(f"embedding_dim: {a.get('embedding_dim')}")
    d.add_paragraph(f"embedding_source: {a.get('embedding_source')}")
    d.add_paragraph(f"embedding_model: {a.get('embedding_model')}")

    d.add_heading("Отчёт LLM", level=2)
    add_multiline_report(d, result.llm_report)

    d.add_heading("Числовой вектор текста", level=2)
    d.add_paragraph(f"Размерность: {a.get('embedding_dim')}")
    d.add_paragraph(f"Источник: {a.get('embedding_source')}")
    d.add_paragraph(f"Модель: {a.get('embedding_model') or 'fallback-вектор'}")
    if a.get("embedding_error"):
        d.add_paragraph(f"Примечание: {a.get('embedding_error')}")
    for line in format_numeric_vector(a.get("embedding_vector")):
        d.add_paragraph(line)

    buf = BytesIO()
    d.save(buf)
    buf.seek(0)

    safe_name = doc.filename.replace("/", "_").replace("\\", "_")
    out_name = f"report_{doc_id}_{safe_name}.docx"

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": attachment_content_disposition(out_name)}
    )
@app.get("/analysis_reports/{doc_id}")
def get_reports(doc_id: int):
    """Получить все сохранённые отчёты по документу"""

    db = SessionLocal()

    results = (
        db.query(AnalysisResult)
        .filter(AnalysisResult.doc_id == doc_id)
        .order_by(AnalysisResult.created_at.desc())
        .all()
    )

    db.close()

    return [
        {
            "id": r.id,
            "created_at": serialize_datetime(r.created_at),
            "llm_report": r.llm_report,
            "analysis": {
                "total_words": r.analysis_json.get("total_words") if r.analysis_json else None,
                "unique_words": r.analysis_json.get("unique_words") if r.analysis_json else None,
                "uniqueness": r.analysis_json.get("uniqueness") if r.analysis_json else None,
                "embedding_dim": r.analysis_json.get("embedding_dim") if r.analysis_json else None,
                "embedding_source": r.analysis_json.get("embedding_source") if r.analysis_json else None,
                "embedding_model": r.analysis_json.get("embedding_model") if r.analysis_json else None,
                "selected_checks": r.analysis_json.get("selected_checks") if r.analysis_json else None,
                "selected_check_titles": r.analysis_json.get("selected_check_titles") if r.analysis_json else None,
            }
        }
        for r in results
    ]
